import os
from contextlib import asynccontextmanager
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import docx
from pypdf import PdfReader
from io import BytesIO
from langchain_groq import ChatGroq
from langchain_core.messages import HumanMessage
from langchain.prompts import PromptTemplate
from langchain_community.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.documents import Document
from langchain.chains import create_retrieval_chain
import json
import pymongo
from llmprovider import get_llm
from urllib.parse import quote_plus
from dotenv import load_dotenv
import uuid
import urllib.parse
import numpy as np
import faiss
from transformers import AutoTokenizer, AutoModel
import torch

# Load environment variables from .env file for local testing
# This line should be commented out or removed for production on Render
load_dotenv()

app = FastAPI()

# Configure CORS
origins = ["*"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Global State and Client Initialization ---

# MongoDB Client
mongo_user = os.environ.get("MONGO_USER")
mongo_pass = os.environ.get("MONGO_PASS")

encoded_user = urllib.parse.quote_plus(mongo_user)
encoded_pass = urllib.parse.quote_plus(mongo_pass)

mongo_uri = f"mongodb+srv://{encoded_user}:{encoded_pass}@cluster0.d3fdjg3.mongodb.net/?retryWrites=true&w=majority"
if not mongo_uri:
    raise ValueError("MONGO_URI environment variable not set.")
mongo_client = pymongo.MongoClient(mongo_uri)
db = mongo_client["resume_db"]
resumes_collection = db["resumes"]

# FAISS and Embedding Model
device = "cuda" if torch.cuda.is_available() else "cpu"
tokenizer = AutoTokenizer.from_pretrained("sentence-transformers/all-MiniLM-L6-v2")
model = AutoModel.from_pretrained("sentence-transformers/all-MiniLM-L6-v2").to(device)

def get_embedding(texts):
    """Generates embeddings for a list of texts using the Hugging Face model."""
    inputs = tokenizer(texts, padding=True, truncation=True, return_tensors="pt").to(device)
    with torch.no_grad():
        embeddings = model(**inputs).last_hidden_state.mean(dim=1)
    return embeddings.cpu().numpy()

class VectorIndex:
    def __init__(self, d: int = 384, index_file: str = "faiss_index.bin"):
        self.d = d
        self.index_file = index_file
        self.index = None
        self.doc_ids = []

    def load_index(self):
        """Loads the FAISS index from a file if it exists, otherwise creates a new one."""
        if os.path.exists(self.index_file):
            print(f"Loading existing FAISS index from {self.index_file}...")
            self.index = faiss.read_index(self.index_file)
            # You would also need to load doc_ids, but for this example, we'll keep it simple
            # and assume a new index for each restart. For true persistence, you would
            # also save and load the doc_ids.
            print("FAISS index loaded.")
        else:
            print("Creating new FAISS index.")
            self.index = faiss.IndexFlatL2(self.d)

    def save_index(self):
        """Saves the FAISS index to a file."""
        print(f"Saving FAISS index to {self.index_file}...")
        faiss.write_index(self.index, self.index_file)
        print("FAISS index saved.")

    def add_vectors(self, vectors: np.ndarray, doc_ids: list):
        """Adds vectors and their corresponding document IDs to the index."""
        self.index.add(vectors)
        self.doc_ids.extend(doc_ids)

    def search_vectors(self, query_vectors: np.ndarray, k: int = 5):
        """Performs a search for the top-k nearest neighbors."""
        distances, indices = self.index.search(query_vectors, k)
        return distances, [self.doc_ids[i] for i in indices[0]]

# Initialize FAISS index globally
vector_index = VectorIndex(d=384, index_file="faiss_index.bin")

# --- Lifespan Context Manager ---
@asynccontextmanager
async def lifespan(app: FastAPI):
    # MongoDB connection check
    try:
        mongo_client.admin.command('ping')
        print("MongoDB connection successful!")
    except Exception as e:
        print(f"MongoDB connection failed: {e}")
        raise RuntimeError("MongoDB connection is required for the application to function.")

    # Load FAISS index on startup
    vector_index.load_index()

    yield
    mongo_client.close()

# --- Helper Functions ---
def extract_text_from_file(file: UploadFile) -> str:
    """Extracts text from a given file based on its content type."""
    content_type = file.content_type
    
    if content_type == 'text/plain':
        return file.file.read().decode('utf-8')
    elif content_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
        doc = docx.Document(file.file)
        return '\n'.join([para.text for para in doc.paragraphs])
    elif content_type == 'application/pdf':
        reader = PdfReader(file.file)
        return ''.join(page.extract_text() for page in reader.pages)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

def extract_data_with_llm(resume_text: str):
    """Extracts structured data from resume text using an LLM."""
    llm = get_llm()
    system_prompt = "You are a world-class resume parser and data extractor. You will be given the text content of a resume. Your task is to extract key information and return a JSON object. Ensure all fields are present, even if empty. The JSON object MUST have the following keys: 'name', 'contact_info' (an object with 'email', 'phone', 'linkedin'), 'summary', 'education' (an array of objects with 'degree', 'institution', 'year'), 'experience' (an array of objects with 'title', 'company', 'start_date', 'end_date', 'description'), and 'skills' (an array of strings). Do not include any other text, just the JSON."
    
    messages = [
        HumanMessage(content=f"{system_prompt}\n\nResume Text:\n{resume_text}")
    ]
    response = llm.invoke(messages)
    return json.loads(response.content)

def generate_markdown_with_llm(extracted_data: dict):
    """Generates a professional markdown resume from extracted data."""
    llm = get_llm()
    prompt = f"""You are a professional resume writer. Rewrite the following resume data into a well-formatted and professional markdown resume. Use markdown headings, bullet points, and bold text to make it clean and readable. Do not include any extra sentences or conversational text outside of the resume content itself.
    
    Resume Data:\n{json.dumps(extracted_data, indent=2)}
    
    Markdown Resume:
    """
    response = llm.invoke([HumanMessage(content=prompt)])
    return response.content

# --- API Endpoints ---
@app.post("/api/process")
async def process_resume_endpoint(resume_file: UploadFile = File(...), user_id: str = Form(...)):
    """API endpoint to process a resume file."""
    try:
        resume_text = extract_text_from_file(resume_file)
        extracted_data = extract_data_with_llm(resume_text)
        markdown_content = generate_markdown_with_llm(extracted_data)
        
        # Save to MongoDB
        resume_doc = {
            "user_id": user_id,
            "extracted_data": extracted_data,
            "markdown_content": markdown_content,
            "original_text": resume_text,
            "uploaded_at": str(uuid.uuid4())
        }
        result = resumes_collection.insert_one(resume_doc)
        resume_id = str(result.inserted_id)
        
        # Save to FAISS for vector search
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        chunks = text_splitter.split_text(resume_text)
        
        chunk_embeddings = get_embedding(chunks)
        chunk_ids = [f"{resume_id}_{i}" for i in range(len(chunks))]
        
        vector_index.add_vectors(chunk_embeddings, chunk_ids)
        vector_index.save_index() # Save the index after adding data
        
        return {
            "message": "Resume processed successfully",
            "extracted_data": extracted_data,
            "markdown_content": markdown_content,
            "resume_id": resume_id
        }
    except Exception as e:
        print(e)
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {e}")

@app.post("/api/chat")
async def chat_endpoint(query: dict):
    """API endpoint for chat and Q&A using RAG."""
    user_query = query.get("query", "")
    
    if vector_index.index is None or vector_index.index.ntotal == 0:
        return {"response": "No resumes have been processed yet. Please upload a resume first."}
        
    # Get embedding for the query
    query_embedding = get_embedding([user_query])

    # Query FAISS for relevant documents
    distances, matching_ids = vector_index.search_vectors(query_embedding, k=5)
    
    # Retrieve the full resume documents from MongoDB based on chunk IDs
    retrieved_docs_text = []
    
    # Find unique resume IDs from the matched chunks
    unique_resume_ids = list(set([_id.split('_')[0] for _id in matching_ids]))
    
    for resume_id in unique_resume_ids:
        doc = resumes_collection.find_one({"_id": pymongo.ObjectId(resume_id)})
        if doc:
            retrieved_docs_text.append(doc['original_text'])

    retrieved_docs = [Document(page_content=text) for text in retrieved_docs_text]

    llm = get_llm()
    prompt = PromptTemplate.from_template("""
    You are a helpful assistant who answers questions about resumes.
    Use the following context to answer the user's question.
    If you don't know the answer, just say you don't have enough information.
    Context: {context}
    Question: {input}
    """)
    
    document_chain = create_stuff_documents_chain(llm, prompt)
    response = document_chain.invoke({
        "input": user_query,
        "context": retrieved_docs
    })
    
    return {"response": response}

@app.get("/api/resumes/list")
async def list_resumes_endpoint():
    """API endpoint to list all processed resumes."""
    resumes_list = list(resumes_collection.find({}, {"_id": 1, "extracted_data.name": 1, "extracted_data.skills": 1}))
    # Convert ObjectId to string for JSON serialization
    for resume in resumes_list:
        resume["_id"] = str(resume["_id"])
    return {"resumes": resumes_list}

@app.post("/api/resumes/filter")
async def filter_resumes_endpoint(query: dict):
    """API endpoint to filter resumes based on a query using the LLM."""
    user_query = query.get("query", "")
    
    # Get embedding for the query
    query_embedding = get_embedding([user_query])
    
    if vector_index.index is None or vector_index.index.ntotal == 0:
        return {"resumes": []}
    
    # Query FAISS for relevant documents
    distances, matching_chunk_ids = vector_index.search_vectors(query_embedding, k=10)
    
    # Get unique resume IDs from the search results
    matching_ids = {chunk_id.split('_')[0] for chunk_id in matching_chunk_ids}
    
    # Retrieve the full resumes from MongoDB
    filtered_resumes = list(resumes_collection.find({"_id": {"$in": [pymongo.ObjectId(id) for id in matching_ids]}}))
    
    # Convert ObjectId to string for JSON serialization
    for resume in filtered_resumes:
        resume["_id"] = str(resume["_id"])

    # Use the LLM to refine the list and select the best matches
    llm = get_llm()
    
    # Prepare the data for the LLM
    candidate_resumes_string = "\n---\n".join([
        f"ID: {r['_id']}\nName: {r['extracted_data']['name']}\nSkills: {', '.join(r['extracted_data']['skills'])}\nSummary: {r['extracted_data']['summary']}"
        for r in filtered_resumes
    ])
    
    prompt = f"""
    Based on the following list of resumes, identify and list the IDs of the resumes that best match the query.
    Return only a comma-separated list of IDs. Do not include any other text, explanations, or formatting.
    If no resumes match, return an empty string.

    Resumes:
    {candidate_resumes_string}

    Query: {user_query}

    Matching IDs:
    """
    
    response = llm.invoke([HumanMessage(content=prompt)])
    matching_ids_str = response.content.strip()
    
    if not matching_ids_str:
        return {"resumes": []}
    
    final_matching_ids = [id.strip() for id in matching_ids_str.split(',') if id.strip()]
    
    # Fetch the final list of full resumes from the initial filtered set
    final_filtered_resumes = [r for r in filtered_resumes if r["_id"] in final_matching_ids]

    return {"resumes": final_filtered_resumes}
