import os
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import docx
from pypdf import PdfReader
from io import BytesIO
from langchain_groq import ChatGroq
from langchain_core.messages import HumanMessage
from langchain.prompts import PromptTemplate
from langchain_community.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OllamaEmbeddings
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.documents import Document
from langchain.chains import create_retrieval_chain
from langchain_community.vectorstores import FAISS
import json
from llm_provider import get_llm

app = FastAPI()

# In-memory storage for resumes and vector store
in_memory_db = {}
vector_store = None

# Configure CORS
origins = ["*"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def extract_text_from_file(file: UploadFile) -> str:
    """Extracts text from a given file based on its content type."""
    content_type = file.content_type
    
    if content_type == 'text/plain':
        return file.file.read().decode('utf-8')
    elif content_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
        doc = docx.Document(file.file)
        return '\n'.join([para.text for para in doc.paragraphs])
    elif content_type == 'application/pdf':
        reader = PdfReader(file.file)
        return ''.join(page.extract_text() for page in reader.pages)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

def extract_data_with_llm(resume_text: str):
    """Extracts structured data from resume text using an LLM."""
    llm = get_llm()
    system_prompt = "You are a world-class resume parser and data extractor. You will be given the text content of a resume. Your task is to extract key information and return a JSON object. Ensure all fields are present, even if empty. The JSON object MUST have the following keys: 'name', 'contact_info' (an object with 'email', 'phone', 'linkedin'), 'summary', 'education' (an array of objects with 'degree', 'institution', 'year'), 'experience' (an array of objects with 'title', 'company', 'start_date', 'end_date', 'description'), and 'skills' (an array of strings). Do not include any other text, just the JSON."
    
    messages = [
        HumanMessage(content=f"{system_prompt}\n\nResume Text:\n{resume_text}")
    ]
    response = llm.invoke(messages)
    return json.loads(response.content)

def generate_markdown_with_llm(extracted_data: dict):
    """Generates a professional markdown resume from extracted data."""
    llm = get_llm()
    prompt = f"""You are a professional resume writer. Rewrite the following resume data into a well-formatted and professional markdown resume. Use markdown headings, bullet points, and bold text to make it clean and readable. Do not include any extra sentences or conversational text outside of the resume content itself.
    
    Resume Data:\n{json.dumps(extracted_data, indent=2)}
    
    Markdown Resume:
    """
    response = llm.invoke([HumanMessage(content=prompt)])
    return response.content

def update_vector_store(text_content, resume_id):
    """Updates the in-memory vector store with the new resume text."""
    global vector_store
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    docs = [Document(page_content=text, metadata={"source": resume_id})]
    split_documents = text_splitter.split_documents(docs)
    embeddings = OllamaEmbeddings(model="llama3:8b")
    
    if vector_store:
        vector_store.add_documents(split_documents)
    else:
        vector_store = FAISS.from_documents(split_documents, embeddings)

def rag_query(query: str):
    """Performs a RAG query against the in-memory vector store."""
    if not vector_store:
        return "No resumes have been processed yet. Please upload a resume first."
    
    llm = get_llm()
    retriever = vector_store.as_retriever()
    
    prompt = PromptTemplate.from_template("""
    You are a helpful assistant who answers questions about resumes.
    Use the following context to answer the user's question.
    If you don't know the answer, just say you don't have enough information.
    Context: {context}
    Question: {input}
    """)
    
    document_chain = create_stuff_documents_chain(llm, prompt)
    retrieval_chain = create_retrieval_chain(retriever, document_chain)
    
    response = retrieval_chain.invoke({"input": query})
    return response['answer']

@app.post("/api/process")
async def process_resume_endpoint(resume_file: UploadFile = File(...), user_id: str = Form(...)):
    """API endpoint to process a resume file."""
    try:
        resume_text = extract_text_from_file(resume_file)
        extracted_data = extract_data_with_llm(resume_text)
        markdown_content = generate_markdown_with_llm(extracted_data)
        
        # In-memory storage with a unique ID
        resume_id = f"resume_{len(in_memory_db) + 1}"
        in_memory_db[resume_id] = {
            "id": resume_id,
            "user_id": user_id,
            "extracted_data": extracted_data,
            "markdown_content": markdown_content,
            "original_text": resume_text,
        }
        
        # Update RAG model
        update_vector_store(resume_text, resume_id)
        
        return {
            "message": "Resume processed successfully",
            "extracted_data": extracted_data,
            "markdown_content": markdown_content,
            "resume_id": resume_id
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/chat")
async def chat_endpoint(query: dict):
    """API endpoint for chat and Q&A."""
    user_query = query.get("query", "")
    response_text = rag_query(user_query)
    return {"response": response_text}

@app.get("/api/resumes/list")
async def list_resumes_endpoint():
    """API endpoint to list all processed resumes."""
    resumes_list = list(in_memory_db.values())
    return {"resumes": resumes_list}

@app.post("/api/resumes/filter")
async def filter_resumes_endpoint(query: dict):
    """API endpoint to filter resumes based on a query using the LLM."""
    user_query = query.get("query", "")
    
    llm = get_llm()
    
    # Generate a list of all resumes as a single string for the prompt
    resumes_string = "\n---\n".join([
        f"ID: {r['id']}\nName: {r['extracted_data']['name']}\nSkills: {', '.join(r['extracted_data']['skills'])}\nSummary: {r['extracted_data']['summary']}"
        for r in in_memory_db.values()
    ])
    
    prompt = f"""
    Based on the following list of resumes, identify and list the IDs of the resumes that best match the query.
    Return only a comma-separated list of IDs. Do not include any other text, explanations, or formatting.
    If no resumes match, return an empty string.

    Resumes:
    {resumes_string}

    Query: {user_query}

    Matching IDs:
    """
    
    response = llm.invoke([HumanMessage(content=prompt)])
    matching_ids_str = response.content.strip()
    
    if not matching_ids_str:
        return {"resumes": []}
    
    matching_ids = [id.strip() for id in matching_ids_str.split(',') if id.strip()]
    
    # Retrieve the full resume objects for the matching IDs
    filtered_resumes = [in_memory_db[id] for id in matching_ids if id in in_memory_db]
    
    return {"resumes": filtered_resumes}
